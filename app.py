from flask import Flask, request, send_file, jsonify
import pandas as pd
from docx import Document
from datetime import datetime
import zipfile
import tempfile
import os
import io

app = Flask(__name__)

def normalizar_texto(txt):
    return str(txt).strip().lower() if pd.notna(txt) else ""

def primeira_maiuscula(txt):
    return txt.capitalize() if isinstance(txt, str) else ""

def substituir_placeholder_em_paragrafos(paragraphs, marcador, valor):
    for paragraph in paragraphs:
        texto_completo = ''.join(run.text for run in paragraph.runs)
        if marcador in texto_completo:
            novo_texto = texto_completo.replace(marcador, valor)
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = novo_texto

@app.route("/gerar", methods=["POST"])
def gerar():
    try:
        planilhas = request.files.getlist("files[]")
        modelo_file = request.files["modelo"]
        meses = request.form.getlist("meses[]")
        formato = request.form.get("formato")

        consolidado = []
        for i, f in enumerate(planilhas):
            df = pd.read_excel(f)
            mes = meses[i] if i < len(meses) else f"mes{i+1}"
            for _, row in df.iterrows():
                loc = normalizar_texto(row.iloc[3])
                grp = normalizar_texto(row.iloc[8])
                itm = normalizar_texto(row.iloc[11])
                obs = normalizar_texto(row.iloc[27])
                if not itm or not loc:
                    continue
                chave = (loc, itm, grp)
                existente = next((x for x in consolidado if x["chave"] == chave), None)
                if existente:
                    existente["meses"].add(mes)
                    if obs and (not existente["obs"] or len(obs) > len(existente["obs"])):
                        existente["obs"] = obs
                else:
                    consolidado.append({
                        "chave": chave,
                        "localidade": loc,
                        "item": itm,
                        "grupo": grp,
                        "obs": obs,
                        "meses": {mes}
                    })

        agrupado = {}
        for d in consolidado:
            agrupado.setdefault(d["localidade"], []).append(d)

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path_list = []
            for loc, itens in agrupado.items():
                modelo_file.seek(0)
                doc = Document(modelo_file)
                substituir_placeholder_em_paragrafos(doc.paragraphs, "<<localidade>>", loc.title())
                substituir_placeholder_em_paragrafos(doc.paragraphs, "<<datadaemissao>>", datetime.now().strftime("%d de %B de %Y"))
                tabela = doc.tables[0]
                base = tabela.rows[2]
                tabela._tbl.remove(base._tr)
                for item in itens:
                    nova = tabela.add_row().cells
                    nova[0].text = primeira_maiuscula(item["item"])
                    nova[1].text = primeira_maiuscula(item["obs"])
                    nova[2].text = primeira_maiuscula(item["grupo"])
                    for j, mes in enumerate(["jan", "fev", "mar"]):
                        if mes in item["meses"]:
                            nova[3 + j].text = "X"

                nome = f"{loc.replace(' ', '_')}_{'_'.join(meses)}.docx"
                doc_path = os.path.join(tmpdir, nome)
                doc.save(doc_path)
                docx_path_list.append(doc_path)

            zip_path = os.path.join(tmpdir, "relatorios.zip")
            with zipfile.ZipFile(zip_path, "w") as zipf:
                for doc in docx_path_list:
                    zipf.write(doc, arcname=os.path.basename(doc))

            return send_file(zip_path, as_attachment=True)
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

